# ============================================================
# 🏆 AI Resume Screener v2.0 — Streamlit App
# ============================================================
#
# FEATURES:
#   ✅ Paste JD + Upload up to 10 resumes (PDF/DOCX)
#   ✅ AI-powered semantic ranking (Top 3 + Best 1)
#   ✅ Skill Gap Analysis — shows missing skills per resume
#   ✅ Multi-JD Support — compare resumes against multiple JDs
#   ✅ Dark Mode / Light Mode toggle
#   ✅ Side-by-Side resume comparison (pick any 2)
#   ✅ Minimum score threshold filter (auto-reject)
#
# RUN:  streamlit run app.py
# ============================================================

import streamlit as st
import os
import re
import numpy as np
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity


# ============================================================
# PAGE CONFIG
# ============================================================
st.set_page_config(
    page_title="AI Resume Screener v2.0",
    page_icon="🏆",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ============================================================
# DARK MODE / LIGHT MODE TOGGLE (Sidebar)
# ============================================================
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False

with st.sidebar:
    st.markdown("## ⚙️ Settings")
    dark_mode = st.toggle("🌙 Dark Mode", value=st.session_state.dark_mode)
    st.session_state.dark_mode = dark_mode

# Apply theme CSS
if st.session_state.dark_mode:
    st.markdown(
        """
    <style>
        /* Dark mode overrides */
        .stApp { background-color: #1a1a2e; color: #e0e0e0; }
        .block-container { color: #e0e0e0; }

        .stTextArea > div > div > textarea {
            background-color: #16213e; color: #e0e0e0; border: 1px solid #0f3460;
            border-radius: 10px; font-size: 14px;
        }
        .stFileUploader > div { background-color: #16213e; border-radius: 10px; }

        .stDataFrame { border-radius: 10px; }

        h1, h2, h3, h4, p, span, label, .stMarkdown { color: #e0e0e0 !important; }

        .stButton > button {
            font-size: 18px !important; padding: 12px !important;
            border-radius: 10px !important; font-weight: 600 !important;
        }

        div[data-testid="stMetricValue"] { color: #e0e0e0 !important; }
        div[data-testid="stMetricLabel"] { color: #a0a0a0 !important; }

        .dark-card {
            background: #16213e !important; border-color: #0f3460 !important;
        }
    </style>
    """,
        unsafe_allow_html=True,
    )
else:
    st.markdown(
        """
    <style>
        .block-container { padding-top: 2rem; }
        #MainMenu { visibility: hidden; }
        footer { visibility: hidden; }
        .stFileUploader > div > div { border-radius: 10px; }
        .stButton > button {
            font-size: 18px !important; padding: 12px !important;
            border-radius: 10px !important; font-weight: 600 !important;
        }
        .stTextArea > div > div > textarea { border-radius: 10px; font-size: 14px; }
        .stDataFrame { border-radius: 10px; }
    </style>
    """,
        unsafe_allow_html=True,
    )

# Theme-aware colors
BG_CARD = "#16213e" if dark_mode else "#ffffff"
BG_HERO = (
    "linear-gradient(135deg, #0f3460, #16213e)"
    if dark_mode
    else "linear-gradient(135deg, #e8f8f5, #d5f5e3)"
)
TEXT_PRIMARY = "#e0e0e0" if dark_mode else "#1a5276"
TEXT_SECONDARY = "#a0a0a0" if dark_mode else "#7f8c8d"
TEXT_DARK = "#e0e0e0" if dark_mode else "#2c3e50"
BORDER_LIGHT = "#0f3460" if dark_mode else "#e0e0e0"
PROGRESS_BG = "#2c3e50" if dark_mode else "#ecf0f1"


# ============================================================
# LOAD MODEL (cached)
# ============================================================
@st.cache_resource
def load_model():
    """Load sentence transformer model — cached across reruns."""
    return SentenceTransformer("all-MiniLM-L6-v2")


# ============================================================
# TEXT EXTRACTION
# ============================================================
def extract_text_from_docx(file):
    """Extract text from .docx (paragraphs + tables)."""
    try:
        doc = Document(file)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception as e:
        st.error(f"❌ Error reading DOCX: {e}")
        return ""


def extract_text_from_pdf(file):
    """Extract text from .pdf across all pages."""
    try:
        reader = PdfReader(file)
        parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())
        return "\n".join(parts)
    except Exception as e:
        st.error(f"❌ Error reading PDF: {e}")
        return ""


def extract_text(uploaded_file):
    """Route to correct parser based on extension."""
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    elif name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    return ""


# ============================================================
# SKILL GAP ANALYSIS
# ============================================================
def extract_skills_from_jd(jd_text):
    """
    Extract skills/keywords from job description.
    Looks for common tech skills, tools, and requirement patterns.
    """
    # Common tech skills & tools dictionary
    known_skills = [
        "python", "java", "javascript", "typescript", "go", "rust", "c\\+\\+", "c#",
        "r\\b", "scala", "ruby", "php", "swift", "kotlin",
        "pytorch", "tensorflow", "keras", "scikit-learn", "sklearn",
        "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
        "spark", "hadoop", "airflow", "kafka", "flink",
        "sql", "nosql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
        "docker", "kubernetes", "k8s", "aws", "azure", "gcp", "cloud",
        "sagemaker", "mlflow", "kubeflow", "mlops", "ci/cd", "ci cd",
        "git", "github", "gitlab", "jenkins", "terraform",
        "react", "node\\.js", "nodejs", "django", "flask", "fastapi",
        "nlp", "natural language processing", "computer vision",
        "bert", "gpt", "transformers", "llm", "large language model",
        "rag", "retrieval augmented generation", "langchain",
        "cnn", "rnn", "lstm", "gan",
        "deep learning", "machine learning", "reinforcement learning",
        "statistics", "probability", "linear algebra",
        "a/b testing", "ab testing", "hypothesis testing",
        "data pipeline", "etl", "data engineering",
        "tableau", "power bi", "looker",
        "agile", "scrum", "jira",
        "pinecone", "weaviate", "chromadb", "vector database",
        "onnx", "tensorrt", "triton",
        "rest api", "graphql", "microservices",
        "xgboost", "random forest", "gradient boosting",
        "object detection", "image classification",
        "time series", "forecasting", "recommendation system",
        "excel", "powerpoint", "word",
        "communication", "leadership", "teamwork", "management",
        "seo", "sem", "google analytics", "hubspot", "crm",
        "content marketing", "social media", "email marketing",
        "figma", "photoshop", "illustrator",
    ]

    jd_lower = jd_text.lower()
    found_skills = []

    for skill in known_skills:
        pattern = r"\b" + skill + r"\b"
        if re.search(pattern, jd_lower):
            # Clean up the skill name for display
            display_name = skill.replace("\\b", "").replace("\\+\\+", "++").replace("\\.", ".")
            found_skills.append(display_name)

    # Remove duplicates and sort
    found_skills = sorted(set(found_skills))
    return found_skills


def analyze_skill_gaps(jd_skills, resume_text):
    """
    Check which JD skills are present/missing in a resume.
    Returns: (matched_skills, missing_skills)
    """
    resume_lower = resume_text.lower()
    matched = []
    missing = []

    for skill in jd_skills:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, resume_lower):
            matched.append(skill)
        else:
            missing.append(skill)

    return matched, missing


# ============================================================
# RANKING FUNCTION
# ============================================================
def rank_resumes(jd_text, resume_data, model):
    """
    Compute semantic similarity between JD and each resume.
    Returns sorted DataFrame with Rank, Filename, Score_Pct.
    """
    jd_embedding = model.encode([jd_text.strip()])
    resume_names = list(resume_data.keys())
    resume_texts = list(resume_data.values())
    resume_embeddings = model.encode(resume_texts, batch_size=8)
    similarities = cosine_similarity(jd_embedding, resume_embeddings)[0]

    results = pd.DataFrame({"Filename": resume_names, "Score": similarities})
    results = results.sort_values("Score", ascending=False).reset_index(drop=True)
    results["Rank"] = range(1, len(results) + 1)
    results["Score_Pct"] = (results["Score"] * 100).round(1)
    return results


# ============================================================
# SIDEBAR: Settings & Filters
# ============================================================
with st.sidebar:
    st.markdown("---")
    st.markdown("## 🎚️ Filters")

    min_score = st.slider(
        "Minimum Match Score (%)",
        min_value=0,
        max_value=100,
        value=0,
        step=5,
        help="Resumes scoring below this threshold will be auto-rejected.",
    )

    if min_score > 0:
        st.info(f"🚫 Resumes below **{min_score}%** will be rejected")

    st.markdown("---")
    st.markdown("## ℹ️ About")
    st.caption(
        "**AI Resume Screener v2.0**\n\n"
        "Model: all-MiniLM-L6-v2\n\n"
        "Embedding Dims: 384\n\n"
        "Accepts: .pdf & .docx"
    )


# ============================================================
# HEADER
# ============================================================
st.markdown(
    f"""
<div style="text-align:center; padding: 10px 0 5px 0;">
    <h1 style="margin:0; font-size:2.5rem; color:{TEXT_PRIMARY};">🏆 AI Resume Screener <sup style="font-size:14px; color:{TEXT_SECONDARY};">v2.0</sup></h1>
    <p style="color:{TEXT_SECONDARY}; font-size:17px; margin-top:5px;">
        Paste Job Description &nbsp;·&nbsp; Upload Resumes &nbsp;·&nbsp; Find the Best Match
    </p>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown("---")


# ============================================================
# MULTI-JD SUPPORT: Tabs for single vs multi JD
# ============================================================
mode_tab1, mode_tab2 = st.tabs(["📋 Single Job Description", "📋📋 Compare Multiple JDs"])


# ==============================================================
# SHARED: Resume upload (used by both tabs)
# ==============================================================
def render_resume_upload(key_suffix=""):
    """Render the resume upload widget. Returns uploaded files."""
    st.markdown("### 📄 Upload Resumes")
    st.caption("Upload up to 10 resumes · .pdf and .docx only")

    files = st.file_uploader(
        "Upload Resumes",
        type=["pdf", "docx"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key=f"uploader_{key_suffix}",
    )

    if files:
        if len(files) > 10:
            st.warning("⚠️ Max 10. Only first 10 processed.")
            files = files[:10]
        st.success(f"✅ {len(files)} resume(s) uploaded")
        for f in files:
            icon = "📕 PDF" if f.name.lower().endswith(".pdf") else "📄 DOCX"
            st.caption(f"{icon}  {f.name}  ({f.size / 1024:.0f} KB)")
    else:
        st.info("⬆️ Drag & drop your resume files here")

    return files


def parse_resumes(uploaded_files):
    """Extract text from all uploaded files. Returns dict {name: text}."""
    data = {}
    errors = []
    for f in uploaded_files:
        f.seek(0)
        text = extract_text(f)
        if text and text.strip():
            data[f.name] = text.strip()
        else:
            errors.append(f.name)
    if errors:
        st.warning(f"⚠️ Could not extract text from: {', '.join(errors)}")
    return data


# ============================================================
# DISPLAY HELPERS
# ============================================================
def display_results(results, resume_data, jd_text, min_score_threshold, model):
    """Render full results: hero banner, top 3 cards, table, chart, skill gaps, comparison."""

    # --- Apply minimum score filter ---
    passed = results[results["Score_Pct"] >= min_score_threshold].copy()
    rejected = results[results["Score_Pct"] < min_score_threshold].copy()

    if not rejected.empty:
        st.markdown(
            f"""
        <div style="padding:12px 20px; background:#fdedec; border-radius:10px; 
                    border:1px solid #e74c3c; margin:10px 0;">
            <p style="margin:0; color:#c0392b; font-size:14px;">
                🚫 <b>{len(rejected)} resume(s) auto-rejected</b> (below {min_score_threshold}% threshold): 
                {', '.join(rejected['Filename'].tolist())}
            </p>
        </div>
        """,
            unsafe_allow_html=True,
        )

    if passed.empty:
        st.error("❌ All resumes were rejected by the minimum score filter. Try lowering the threshold.")
        return

    # Re-rank passed resumes
    passed = passed.reset_index(drop=True)
    passed["Rank"] = range(1, len(passed) + 1)

    # --- Best Match Hero ---
    best = passed.iloc[0]
    st.markdown(
        f"""
    <div style="padding:30px; background:{BG_HERO}; border-radius:15px;
                border:2px solid #27ae60; text-align:center; margin:10px 0 25px 0;
                box-shadow: 0 4px 15px rgba(39,174,96,0.15);">
        <p style="font-size:13px; color:#27ae60; margin:0; font-weight:700;
                  letter-spacing:1px; text-transform:uppercase;">✨ Recommended Candidate</p>
        <h2 style="color:{TEXT_PRIMARY}; margin:10px 0 5px 0; font-size:1.8rem;">{best['Filename']}</h2>
        <p style="font-size:42px; font-weight:800; color:#27ae60; margin:5px 0;">{best['Score_Pct']}%</p>
        <p style="color:{TEXT_SECONDARY}; margin:0; font-size:14px;">
            Highest semantic match out of {len(passed)} qualifying resume(s)
        </p>
    </div>
    """,
        unsafe_allow_html=True,
    )

    # --- Top 3 Cards ---
    st.markdown("### 🏆 Top 3 Candidates")
    top_n = min(3, len(passed))
    top3 = passed.head(top_n)

    medal_icons = ["🥇", "🥈", "🥉"]
    medal_labels = ["BEST MATCH", "2nd Best", "3rd Best"]
    card_bgs = (
        ["#1a3a5c", "#16213e", "#2c2c1a"] if dark_mode else ["#fef9e7", "#f8f9f9", "#fdf2e9"]
    )
    border_cs = ["#f1c40f", "#bdc3c7", "#e67e22"]
    score_cs = ["#27ae60", "#2980b9", "#e67e22"]

    cols = st.columns(top_n)
    for i, (_, row) in enumerate(top3.iterrows()):
        with cols[i]:
            is_best = i == 0
            glow = "box-shadow: 0 4px 20px rgba(241,196,15,0.3);" if is_best else ""
            st.markdown(
                f"""
            <div style="background:{card_bgs[i]}; border:2px solid {border_cs[i]};
                        border-radius:14px; padding:22px 15px; text-align:center;
                        min-height:240px; {glow}">
                <div style="font-size:44px; margin-bottom:5px;">{medal_icons[i]}</div>
                <p style="font-size:11px; color:{score_cs[i]}; font-weight:700;
                          margin:5px 0; letter-spacing:0.5px; text-transform:uppercase;">
                    {medal_labels[i]}</p>
                <p style="font-size:14px; font-weight:700; color:{TEXT_DARK};
                          margin:10px 0; word-wrap:break-word; line-height:1.3;">
                    {row['Filename']}</p>
                <p style="font-size:34px; font-weight:800; color:{score_cs[i]}; margin:8px 0;">
                    {row['Score_Pct']}%</p>
                <div style="margin-top:12px; background:{PROGRESS_BG}; border-radius:8px;
                            overflow:hidden; height:8px;">
                    <div style="width:{row['Score_Pct']}%; height:100%;
                                background:{score_cs[i]}; border-radius:8px;"></div>
                </div>
            </div>
            """,
                unsafe_allow_html=True,
            )

    st.markdown("")

    # --- Full Ranking Table ---
    st.markdown("---")
    st.markdown("### 📊 Complete Ranking")

    tbl = passed[["Rank", "Filename", "Score_Pct"]].copy()
    tbl.columns = ["Rank", "Resume", "Match %"]
    tbl["File Type"] = tbl["Resume"].apply(
        lambda x: "📕 PDF" if x.lower().endswith(".pdf") else "📄 DOCX"
    )
    tbl = tbl[["Rank", "Resume", "File Type", "Match %"]]

    st.dataframe(
        tbl,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Rank": st.column_config.NumberColumn("🏅 Rank", width="small"),
            "Resume": st.column_config.TextColumn("📄 Resume", width="large"),
            "File Type": st.column_config.TextColumn("Type", width="small"),
            "Match %": st.column_config.ProgressColumn(
                "Match Score", min_value=0, max_value=100, format="%.1f%%"
            ),
        },
    )

    # --- Bar Chart ---
    st.markdown("### 📈 Visual Comparison")
    chart = passed[["Filename", "Score_Pct"]].copy()
    chart.columns = ["Resume", "Match %"]
    chart = chart.sort_values("Match %", ascending=True)
    st.bar_chart(chart.set_index("Resume"), horizontal=True, color="#3498db")

    # --- Summary Metrics ---
    st.markdown("---")
    mcols = st.columns(4)
    with mcols[0]:
        st.metric("📄 Analyzed", len(passed))
    with mcols[1]:
        st.metric("🏆 Best Score", f"{passed.iloc[0]['Score_Pct']}%")
    with mcols[2]:
        st.metric("📉 Lowest Score", f"{passed.iloc[-1]['Score_Pct']}%")
    with mcols[3]:
        st.metric("📊 Average", f"{passed['Score_Pct'].mean():.1f}%")

    if not rejected.empty:
        st.metric("🚫 Rejected", len(rejected))

    # ==========================================
    # SKILL GAP ANALYSIS
    # ==========================================
    st.markdown("---")
    st.markdown("### 🔍 Skill Gap Analysis")
    st.caption("Shows which job description skills each resume has or is missing.")

    jd_skills = extract_skills_from_jd(jd_text)

    if not jd_skills:
        st.info("ℹ️ No specific skills detected in the job description.")
    else:
        st.markdown(
            f"**{len(jd_skills)} skills detected in JD:** {', '.join(jd_skills)}"
        )
        st.markdown("")

        for _, row in passed.iterrows():
            fname = row["Filename"]
            resume_text = resume_data.get(fname, "")
            matched, missing = analyze_skill_gaps(jd_skills, resume_text)
            match_pct = (
                round(len(matched) / len(jd_skills) * 100) if jd_skills else 0
            )

            skill_border = "#27ae60" if match_pct >= 70 else ("#f39c12" if match_pct >= 40 else "#e74c3c")

            with st.expander(
                f"{'🟢' if match_pct >= 70 else ('🟡' if match_pct >= 40 else '🔴')} "
                f"{fname} — {match_pct}% skills matched "
                f"({len(matched)}/{len(jd_skills)})"
            ):
                col_m, col_x = st.columns(2)
                with col_m:
                    st.markdown("**✅ Skills Found:**")
                    if matched:
                        st.success(", ".join(matched))
                    else:
                        st.caption("None")
                with col_x:
                    st.markdown("**❌ Skills Missing:**")
                    if missing:
                        st.error(", ".join(missing))
                    else:
                        st.caption("None — perfect match!")

                # Mini progress bar
                st.progress(match_pct / 100, text=f"Skill coverage: {match_pct}%")

    # ==========================================
    # SIDE-BY-SIDE COMPARISON
    # ==========================================
    if len(passed) >= 2:
        st.markdown("---")
        st.markdown("### ⚖️ Side-by-Side Resume Comparison")
        st.caption("Select any 2 resumes to compare them head-to-head.")

        compare_options = passed["Filename"].tolist()
        comp_cols = st.columns(2)

        with comp_cols[0]:
            res_a = st.selectbox("Resume A", compare_options, index=0, key="comp_a")
        with comp_cols[1]:
            default_b = 1 if len(compare_options) > 1 else 0
            res_b = st.selectbox("Resume B", compare_options, index=default_b, key="comp_b")

        if res_a and res_b:
            row_a = passed[passed["Filename"] == res_a].iloc[0]
            row_b = passed[passed["Filename"] == res_b].iloc[0]

            text_a = resume_data.get(res_a, "")
            text_b = resume_data.get(res_b, "")

            words_a = len(text_a.split())
            words_b = len(text_b.split())

            # Skill gaps for both
            matched_a, missing_a = analyze_skill_gaps(jd_skills, text_a) if jd_skills else ([], [])
            matched_b, missing_b = analyze_skill_gaps(jd_skills, text_b) if jd_skills else ([], [])
            skill_pct_a = round(len(matched_a) / len(jd_skills) * 100) if jd_skills else 0
            skill_pct_b = round(len(matched_b) / len(jd_skills) * 100) if jd_skills else 0

            # Winner highlight
            winner_color_a = "#27ae60" if row_a["Score_Pct"] >= row_b["Score_Pct"] else TEXT_SECONDARY
            winner_color_b = "#27ae60" if row_b["Score_Pct"] >= row_a["Score_Pct"] else TEXT_SECONDARY

            c1, c2 = st.columns(2)

            with c1:
                winner_a = " 👑" if row_a["Score_Pct"] > row_b["Score_Pct"] else ""
                st.markdown(
                    f"""
                <div style="background:{BG_CARD}; border:2px solid {winner_color_a};
                            border-radius:12px; padding:20px; text-align:center;">
                    <h3 style="color:{TEXT_PRIMARY}; margin:0;">{res_a}{winner_a}</h3>
                    <p style="font-size:36px; font-weight:800; color:{winner_color_a}; margin:10px 0;">
                        {row_a['Score_Pct']}%</p>
                    <p style="color:{TEXT_SECONDARY};">Rank #{int(row_a['Rank'])} · {words_a} words · Skills: {skill_pct_a}%</p>
                </div>
                """,
                    unsafe_allow_html=True,
                )
                if jd_skills:
                    st.caption(f"✅ Has: {', '.join(matched_a[:8])}{'...' if len(matched_a) > 8 else ''}")
                    st.caption(f"❌ Missing: {', '.join(missing_a[:8])}{'...' if len(missing_a) > 8 else ''}")

            with c2:
                winner_b = " 👑" if row_b["Score_Pct"] > row_a["Score_Pct"] else ""
                st.markdown(
                    f"""
                <div style="background:{BG_CARD}; border:2px solid {winner_color_b};
                            border-radius:12px; padding:20px; text-align:center;">
                    <h3 style="color:{TEXT_PRIMARY}; margin:0;">{res_b}{winner_b}</h3>
                    <p style="font-size:36px; font-weight:800; color:{winner_color_b}; margin:10px 0;">
                        {row_b['Score_Pct']}%</p>
                    <p style="color:{TEXT_SECONDARY};">Rank #{int(row_b['Rank'])} · {words_b} words · Skills: {skill_pct_b}%</p>
                </div>
                """,
                    unsafe_allow_html=True,
                )
                if jd_skills:
                    st.caption(f"✅ Has: {', '.join(matched_b[:8])}{'...' if len(matched_b) > 8 else ''}")
                    st.caption(f"❌ Missing: {', '.join(missing_b[:8])}{'...' if len(missing_b) > 8 else ''}")

            # Comparison table
            comp_df = pd.DataFrame(
                {
                    "Metric": ["Match Score", "Rank", "Word Count", "Skill Coverage", "Missing Skills"],
                    res_a: [
                        f"{row_a['Score_Pct']}%",
                        f"#{int(row_a['Rank'])}",
                        str(words_a),
                        f"{skill_pct_a}%",
                        str(len(missing_a)),
                    ],
                    res_b: [
                        f"{row_b['Score_Pct']}%",
                        f"#{int(row_b['Rank'])}",
                        str(words_b),
                        f"{skill_pct_b}%",
                        str(len(missing_b)),
                    ],
                }
            )
            st.markdown("")
            st.dataframe(comp_df, use_container_width=True, hide_index=True)


# ============================================================
# TAB 1: SINGLE JD MODE
# ============================================================
with mode_tab1:
    col_jd, col_div, col_res = st.columns([5, 0.2, 5])

    with col_jd:
        st.markdown("### 📋 Job Description")
        st.caption("Paste the complete job description below")
        jd_text_single = st.text_area(
            "JD",
            height=420,
            placeholder="Paste the full job description here...\n\nExample:\nJob Title: Senior ML Engineer\nRequirements:\n- 4+ years in ML...",
            label_visibility="collapsed",
            key="jd_single",
        )
        if jd_text_single and jd_text_single.strip():
            st.success(f"✅ JD loaded — {len(jd_text_single.split())} words")
        else:
            st.warning("⬆️ Paste your job description above")

    with col_div:
        st.markdown(
            f'<div style="border-left:2px solid {BORDER_LIGHT}; height:480px; margin:0 auto;"></div>',
            unsafe_allow_html=True,
        )

    with col_res:
        uploaded_single = render_resume_upload("single")

    st.markdown("---")

    can_rank_single = bool(
        jd_text_single and jd_text_single.strip() and uploaded_single and len(uploaded_single) > 0
    )

    if st.button(
        "🚀  Rank Resumes — Find Best Match",
        type="primary",
        use_container_width=True,
        disabled=not can_rank_single,
        key="btn_single",
    ):
        with st.spinner("📄 Parsing resumes..."):
            resume_data = parse_resumes(uploaded_single)

        if not resume_data:
            st.error("❌ No text could be extracted. Check your files.")
            st.stop()

        with st.spinner("🧠 AI is analyzing resumes..."):
            model = load_model()
            results = rank_resumes(jd_text_single, resume_data, model)

        display_results(results, resume_data, jd_text_single, min_score, model)

    elif not can_rank_single:
        st.markdown(
            f'<p style="text-align:center; color:{TEXT_SECONDARY}; font-size:14px;">'
            "👆 Paste a job description AND upload at least one resume to get started.</p>",
            unsafe_allow_html=True,
        )


# ============================================================
# TAB 2: MULTI-JD MODE
# ============================================================
with mode_tab2:
    st.markdown("### 📋📋 Compare Resumes Against Multiple Job Descriptions")
    st.caption(
        "Paste up to 3 different job descriptions to see which roles each resume fits best."
    )

    num_jds = st.selectbox("How many JDs to compare?", [2, 3], index=0, key="num_jds")

    jd_cols = st.columns(num_jds)
    jd_texts_multi = []

    for i in range(num_jds):
        with jd_cols[i]:
            st.markdown(f"**JD #{i + 1}**")
            jd_input = st.text_area(
                f"Job Description {i + 1}",
                height=250,
                placeholder=f"Paste job description #{i + 1} here...",
                label_visibility="collapsed",
                key=f"jd_multi_{i}",
            )
            jd_texts_multi.append(jd_input)
            if jd_input and jd_input.strip():
                st.success(f"✅ {len(jd_input.split())} words")

    st.markdown("---")
    uploaded_multi = render_resume_upload("multi")
    st.markdown("---")

    all_jds_filled = all(jd.strip() for jd in jd_texts_multi)
    can_rank_multi = bool(all_jds_filled and uploaded_multi and len(uploaded_multi) > 0)

    if st.button(
        "🚀  Compare Across All JDs",
        type="primary",
        use_container_width=True,
        disabled=not can_rank_multi,
        key="btn_multi",
    ):
        with st.spinner("📄 Parsing resumes..."):
            resume_data_multi = parse_resumes(uploaded_multi)

        if not resume_data_multi:
            st.error("❌ No text extracted.")
            st.stop()

        model = load_model()

        # Rank against each JD
        all_results = {}
        for i, jd_t in enumerate(jd_texts_multi):
            with st.spinner(f"🧠 Analyzing against JD #{i + 1}..."):
                all_results[f"JD #{i + 1}"] = rank_resumes(jd_t, resume_data_multi, model)

        # --- Build comparison matrix ---
        st.markdown("---")
        st.markdown("### 📊 Cross-JD Comparison Matrix")
        st.caption("Each cell shows the match % of a resume against each job description.")

        resume_names_multi = list(resume_data_multi.keys())
        matrix_data = {"Resume": resume_names_multi}

        for jd_label, res_df in all_results.items():
            scores = {}
            for _, row in res_df.iterrows():
                scores[row["Filename"]] = row["Score_Pct"]
            matrix_data[jd_label] = [scores.get(name, 0) for name in resume_names_multi]

        matrix_df = pd.DataFrame(matrix_data)

        # Add "Best Fit" column
        jd_columns = [c for c in matrix_df.columns if c != "Resume"]
        matrix_df["Best Fit"] = matrix_df[jd_columns].idxmax(axis=1)
        matrix_df["Best Score"] = matrix_df[jd_columns].max(axis=1)
        matrix_df = matrix_df.sort_values("Best Score", ascending=False).reset_index(drop=True)

        st.dataframe(
            matrix_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Resume": st.column_config.TextColumn("📄 Resume", width="large"),
                **{
                    col: st.column_config.ProgressColumn(
                        col, min_value=0, max_value=100, format="%.1f%%"
                    )
                    for col in jd_columns
                },
                "Best Fit": st.column_config.TextColumn("🎯 Best Fit", width="small"),
                "Best Score": st.column_config.NumberColumn("Score", format="%.1f%%", width="small"),
            },
        )

        # --- Per-JD results ---
        st.markdown("---")
        for jd_label, res_df in all_results.items():
            with st.expander(f"🏆 {jd_label} — Detailed Ranking"):
                jd_idx = int(jd_label.split("#")[1]) - 1
                display_results(
                    res_df,
                    resume_data_multi,
                    jd_texts_multi[jd_idx],
                    min_score,
                    model,
                )

    elif not can_rank_multi:
        st.markdown(
            f'<p style="text-align:center; color:{TEXT_SECONDARY}; font-size:14px;">'
            "👆 Fill in all job descriptions AND upload resumes to compare.</p>",
            unsafe_allow_html=True,
        )


# ============================================================
# FOOTER
# ============================================================
st.markdown("---")
st.markdown(
    f"""
<div style="text-align:center; padding:10px; color:{TEXT_SECONDARY};">
    <p style="font-size:13px; margin:0;">
        🤖 Powered by <b>Sentence-Transformers</b> (all-MiniLM-L6-v2) &nbsp;|&nbsp;
        Built with <b>Streamlit</b> &nbsp;|&nbsp;
        Accepts <b>.pdf</b> &amp; <b>.docx</b> only
    </p>
    <p style="font-size:11px; margin:5px 0 0 0;">
        v2.0 — Skill Gap Analysis · Multi-JD · Dark Mode · Side-by-Side Comparison · Score Filter
    </p>
</div>
""",
    unsafe_allow_html=True,
)